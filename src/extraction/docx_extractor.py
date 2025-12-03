"""
DOCX text extraction with robust error handling.
"""
import os
from pathlib import Path


class DOCXExtractor:
    """Extract text from DOCX/DOC files."""
    
    def extract(self, filepath: str) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            filepath: Path to DOCX file
            
        Returns:
            Extracted text string (never None)
            
        Raises:
            ValueError: If file doesn't exist or can't be read
        """
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"DOCX file not found: {filepath}")
        
        # Try python-docx first
        text = self._extract_with_python_docx(filepath)
        if text and len(text.strip()) > 50:
            return text
        
        # Fallback to docx2txt
        text = self._extract_with_docx2txt(filepath)
        if text and len(text.strip()) > 50:
            return text
        
        # If both fail, raise clear error
        raise ValueError(
            f"Could not extract text from DOCX: {filepath}\n"
            f"This file may be:\n"
            f"  • Corrupted or password-protected\n"
            f"  • Empty\n"
            f"  • Using an unsupported format\n"
            f"Please save as .txt or try a different file."
        )
    
    def _extract_with_python_docx(self, filepath: str) -> str:
        """Extract using python-docx library."""
        try:
            from docx import Document
            
            doc = Document(filepath)
            text_parts = []
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return "\n".join(text_parts) if text_parts else ""
            
        except ImportError:
            return ""
        except Exception as e:
            print(f"⚠️ python-docx error: {e}")
            return ""
    
    def _extract_with_docx2txt(self, filepath: str) -> str:
        """Extract using docx2txt library."""
        try:
            import docx2txt
            
            text = docx2txt.process(filepath)
            return text if text else ""
            
        except ImportError:
            return ""
        except Exception as e:
            print(f"⚠️ docx2txt error: {e}")
            return ""